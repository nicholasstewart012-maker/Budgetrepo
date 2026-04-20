import os
import re
import json
import csv
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Iterable, Any

from dotenv import load_dotenv
from docx import Document as DocxDocument
from pypdf import PdfReader
import chromadb

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - optional dependency at import time
    load_workbook = None


try:
    from spark_db import (
        delete_document_chunks,
        delete_document_pages,
        finish_ingestion_run,
        get_document_by_source_path,
        get_document_chunk_count,
        init_db,
        insert_chunks,
        insert_document_pages,
        purge_missing_documents,
        start_ingestion_run,
        upsert_document,
    )
except ModuleNotFoundError:
    from backend.spark_db import (
        delete_document_chunks,
        delete_document_pages,
        finish_ingestion_run,
        get_document_by_source_path,
        get_document_chunk_count,
        init_db,
        insert_chunks,
        insert_document_pages,
        purge_missing_documents,
        start_ingestion_run,
        upsert_document,
    )

try:
    from ocr_service import (
        extract_pdf_page_text,
    )
except ModuleNotFoundError:
    from backend.ocr_service import (
        extract_pdf_page_text,
    )

try:
    from embedding_config import call_embedding_api, get_active_embedding_config
except ModuleNotFoundError:
    from backend.embedding_config import call_embedding_api, get_active_embedding_config

load_dotenv(dotenv_path=Path(__file__).parent.parent / ".env")

INTAKE_FOLDER = Path(os.getenv("INTAKE_FOLDER", r"C:\Spark\intake"))
CHROMA_FOLDER = Path(os.getenv("CHROMA_FOLDER", r"C:\Spark\chromadb"))
CHROMA_COLLECTION = os.getenv("CHROMA_COLLECTION", "spark_documents")
HASH_STORE    = CHROMA_FOLDER / f"_file_hashes_{CHROMA_COLLECTION}.json"
EMBED_BATCH_SIZE = int(os.getenv("EMBED_BATCH_SIZE", "4"))
SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf", ".rtf", ".xlsx", ".xlsm", ".csv"}
DISABLED_EXTENSIONS = {".pptx", ".ppt", ".xls", ".ods"}
PDF_OCR_CHAR_THRESHOLD = int(os.getenv("PDF_OCR_CHAR_THRESHOLD", "40"))
PDF_OCR_WORD_THRESHOLD = int(os.getenv("PDF_OCR_WORD_THRESHOLD", "8"))


def _relative_source_path(filepath: Path) -> str:
    try:
        return filepath.relative_to(INTAKE_FOLDER).as_posix()
    except ValueError:
        return filepath.name


def _make_document_id(source_path: str, source_url: str = "", title: str = "") -> str:
    payload = "|".join([source_path or "", source_url or "", title or ""])
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:32]


def classify_vector_eligibility(chunk: dict, seen_text_counts: dict[str, int] | None = None) -> tuple[bool, str | None]:
    """
    Determines if a chunk should be embedded in Chroma based on its content and metadata.
    Returns (is_eligible, skip_reason).
    """
    text = (chunk.get("text") or "").strip()
    token_count = int(chunk.get("token_count") or 0)
    chunk_type = chunk.get("chunk_type") or "body"

    # 1. Empty or near-empty text
    if not text or len(text) < 10:
        return False, "empty_text"

    # 2. Blacklisted chunk types (low-signal)
    low_signal_types = {
        "heading_only",
        "table_of_contents",
        "disclaimer",
        "revision_history",
        "approval_metadata",
        "admin_metadata",
    }
    if chunk_type in low_signal_types:
        return False, f"low_signal_chunk_type:{chunk_type}"

    # 3. Short chunks
    # Exceptions for spreadsheet key-value pairs which are often short but high-value
    is_spreadsheet_kv = chunk_type == "spreadsheet_key_value"
    if token_count < 25 and not is_spreadsheet_kv:
        return False, "short_chunk"

    # Fallback/OCR threshold
    if chunk_type in {"fallback", "ocr"} and token_count < 40:
        return False, "short_fallback_ocr"

    # 4. Repeated boilerplate (Duplicate detection)
    if seen_text_counts is not None:
        # Normalize text for duplication check (remove whitespace and non-alphanumeric)
        norm_text = re.sub(r"\W+", "", text.lower())
        if len(norm_text) > 20:
            count = seen_text_counts.get(norm_text, 0)
            seen_text_counts[norm_text] = count + 1
            # If we've seen this exact long-ish string more than 3 times, it's likely boilerplate
            # (e.g. copyright notice, footer, standard disclaimer not caught by type)
            if count >= 3:
                return False, "duplicate_boilerplate"

    # 5. Default eligible types
    eligible_types = {
        "body",
        "responsibility",
        "ocr",
        "spreadsheet_row_batch",
        "spreadsheet_key_value",
        "spreadsheet_table_summary",
    }
    if chunk_type in eligible_types:
        return True, None

    # Fallback for other types
    if token_count >= 50:
        return True, None

    return False, "unclassified_low_signal"


def _load_manifest_metadata() -> dict[str, dict[str, Any]]:
    manifest: dict[str, dict[str, Any]] = {}

    def _store(row: dict[str, Any]) -> None:
        cleaned = {k: v for k, v in row.items() if v not in {None, ""}}
        if not cleaned:
            return
        for key in (cleaned.get("source_path"), cleaned.get("file_name"), cleaned.get("title")):
            if key:
                manifest[str(key).replace("\\", "/").lower()] = cleaned

    json_path = INTAKE_FOLDER / "manifest.json"
    if json_path.exists():
        try:
            data = json.loads(json_path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                for key, value in data.items():
                    if isinstance(value, dict):
                        row = dict(value)
                        row.setdefault("source_path", key)
                        _store(row)
            elif isinstance(data, list):
                for item in data:
                    if isinstance(item, dict):
                        _store(item)
        except Exception as exc:
            print(f"[Spark Ingest] Failed to read manifest.json: {exc}")

    csv_path = INTAKE_FOLDER / "manifest.csv"
    if csv_path.exists():
        try:
            with open(csv_path, "r", encoding="utf-8-sig", newline="") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    _store(row)
        except Exception as exc:
            print(f"[Spark Ingest] Failed to read manifest.csv: {exc}")

    return manifest


def _lookup_manifest_metadata(filepath: Path, manifest: dict[str, dict[str, Any]]) -> dict[str, Any]:
    rel = _relative_source_path(filepath).replace("\\", "/").lower()
    for key in (rel, filepath.name.lower(), filepath.stem.lower()):
        if key in manifest:
            return dict(manifest[key])
    return {}


def _normalize_manifest_row(filepath: Path, manifest_row: dict[str, Any], file_hash: str) -> dict[str, Any]:
    source_path = _relative_source_path(filepath)
    source_url = (manifest_row.get("source_url") or manifest_row.get("url") or "").strip() or None
    title = (manifest_row.get("title") or filepath.stem.replace("_", " ")).strip()
    file_name = filepath.name
    file_type = filepath.suffix.lower().lstrip(".") or "unknown"
    department = (manifest_row.get("department") or _infer_department(file_name)).strip() or "General"
    application = (manifest_row.get("application") or "Knowledge Assistant").strip() or "Knowledge Assistant"
    status = (manifest_row.get("status") or "active").strip().lower() or "active"
    audience = (manifest_row.get("audience") or "all").strip() or "all"
    is_admin_only = str(manifest_row.get("is_admin_only") or "0").strip().lower() in {"1", "true", "yes"}
    document_id = (manifest_row.get("document_id") or _make_document_id(source_path, source_url or "", title)).strip()

    return {
        "document_id": document_id,
        "source_path": source_path,
        "source_url": source_url,
        "title": title,
        "file_name": file_name,
        "file_type": file_type,
        "file_size": filepath.stat().st_size,
        "file_hash": file_hash,
        "source_fingerprint": _build_source_fingerprint(filepath, file_hash),
        "department": department,
        "application": application,
        "status": status,
        "audience": audience,
        "is_admin_only": int(is_admin_only),
        "effective_date": (manifest_row.get("effective_date") or "").strip() or None,
        "expiration_date": (manifest_row.get("expiration_date") or "").strip() or None,
        "last_modified": datetime.fromtimestamp(filepath.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
        "last_ingested_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "page_count": None,
        "ingestion_status": "pending",
        "ingestion_error": None,
    }


def _token_count(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text or ""))


def _normalized_lines(text: str) -> list[str]:
    return [line.strip() for line in re.split(r"[\r\n]+", (text or "").strip()) if line.strip()]


def _count_sentence_marks(text: str) -> int:
    return len(re.findall(r"[.!?]", text or ""))


def _looks_like_toc(text: str, section_title: str | None = None, page_number: int | None = None) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    if not cleaned:
        return False

    lower = cleaned.lower()
    lines = _normalized_lines(text)
    leader_count = len(re.findall(r"\.{3,}\s*\d+\b", cleaned))
    page_number_lines = sum(1 for line in lines if re.search(r"\b\d+\s*$", line) and not re.search(r"[.!?]$", line))
    section_row_lines = sum(
        1
        for line in lines
        if re.search(r"^[A-Z][A-Za-z0-9 ,&()/\-]{3,}\s+\.{2,}\s*\d+\s*$", line)
        or re.search(r"^[A-Za-z][A-Za-z0-9 ,&()/\-]{3,}\s+\d+\s*$", line)
    )
    toc_hits = sum(1 for term in ("contents", "table of contents", "toc") if term in lower)
    top_of_page = page_number is None or page_number <= 3
    if toc_hits and top_of_page and (leader_count >= 1 or page_number_lines >= 2 or section_row_lines >= 2):
        return True
    if leader_count >= 2 and page_number_lines >= 2:
        return True
    if section_title and "contents" in section_title.lower() and (leader_count >= 1 or page_number_lines >= 2):
        return True
    if top_of_page and toc_hits and section_row_lines >= 1:
        return True
    return False


def _looks_like_disclaimer(text: str, section_title: str | None = None) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    lower = cleaned.lower()
    title = (section_title or "").lower()
    disclaimer_terms = (
        "disclaimer",
        "does not constitute legal advice",
        "policies constitute statements",
        "should be judged in any legal proceeding",
        "not intended to be",
        "without limitation",
        "no warranty",
    )
    if "disclaimer" in lower or "disclaimer" in title:
        return True
    return any(term in lower for term in disclaimer_terms)


def _looks_like_revision_history(text: str, section_title: str | None = None) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    lower = cleaned.lower()
    title = (section_title or "").lower()
    if "revision history" in lower or "revision history" in title:
        return True
    date_count = len(re.findall(
        r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{4}\b",
        lower,
        re.IGNORECASE,
    ))
    version_count = len(re.findall(r"\bv?\d+\.\d+\b", lower, re.IGNORECASE))
    if date_count >= 2 and version_count >= 1:
        return True
    if len(_normalized_lines(text)) >= 3 and (date_count >= 2 or version_count >= 2):
        return True
    return False


def _looks_like_approval_metadata(text: str, section_title: str | None = None) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    if len(cleaned) < 40:
        return False
    lower = cleaned.lower()
    title = (section_title or "").lower()
    metadata_terms = (
        "sponsoring department",
        "executive sponsor",
        "board committee",
        "committee approval date",
        "effective date",
        "policy owner",
        "document owner",
        "review date",
        "last reviewed",
        "approved by",
        "approval date",
        "document number",
        "classification",
        "owner",
    )
    label_hits = sum(1 for term in metadata_terms if term in lower or term in title)
    colon_count = cleaned.count(":")
    sentence_marks = _count_sentence_marks(cleaned)
    if label_hits >= 2 and sentence_marks <= 2:
        return True
    if colon_count >= 3 and sentence_marks <= 1:
        return True
    return False


def _looks_like_responsibility(text: str, section_title: str | None = None) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    lower = cleaned.lower()
    title = (section_title or "").lower()
    if any(term in title for term in ("responsibilities", "roles and responsibilities", "responsibility")):
        return True
    if any(term in lower for term in ("roles and responsibilities", "responsibilities", "responsible for")):
        return True
    return False


def _looks_like_heading_only(text: str, section_title: str | None = None) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    if not cleaned:
        return True
    if _count_sentence_marks(cleaned) > 0:
        return False
    if len(cleaned) > 120:
        return False
    words = cleaned.split()
    if len(words) <= 14 and (_looks_like_heading(cleaned) or cleaned == (section_title or "").strip()):
        return True
    if len(words) <= 8 and len(cleaned) <= 70:
        return True
    return False


def _looks_like_body_chunk(text: str) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    if len(cleaned) < 60:
        return False
    if _count_sentence_marks(cleaned) == 0:
        return False
    return True


def classify_chunk_type(
    text: str,
    section_title: str | None = None,
    page_number: int | None = None,
    *,
    extraction_method: str | None = None,
    has_text_layer: int | bool | None = None,
) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    if not cleaned:
        return "fallback"

    if _looks_like_toc(cleaned, section_title=section_title, page_number=page_number):
        return "table_of_contents"
    if _looks_like_disclaimer(cleaned, section_title=section_title):
        return "disclaimer"
    if _looks_like_revision_history(cleaned, section_title=section_title):
        return "revision_history"
    if _looks_like_approval_metadata(cleaned, section_title=section_title):
        return "approval_metadata"
    if _looks_like_responsibility(cleaned, section_title=section_title):
        return "responsibility"
    if _looks_like_heading_only(cleaned, section_title=section_title):
        return "heading_only"
    if extraction_method == "ocr" or (has_text_layer is not None and not bool(has_text_layer)):
        if cleaned:
            return "ocr"
    if _looks_like_body_chunk(cleaned):
        return "body"
    return "fallback"


def _chunk_quality_flags(
    text: str,
    *,
    chunk_type: str,
    section_title: str | None = None,
    page_number: int | None = None,
    extraction_method: str | None = None,
    has_text_layer: int | bool | None = None,
    ocr_confidence: float | None = None,
) -> str:
    flags: list[str] = []
    if chunk_type and chunk_type != "body":
        flags.append(chunk_type)
    if extraction_method == "ocr":
        flags.append("ocr")
    if has_text_layer is not None and not bool(has_text_layer):
        flags.append("no_text_layer")
    if page_number is not None and page_number <= 3 and chunk_type == "table_of_contents":
        flags.append("early_page")
    if section_title and chunk_type in {"heading_only", "fallback"}:
        flags.append("section_title_only")
    if ocr_confidence is not None:
        try:
            confidence = float(ocr_confidence)
        except (TypeError, ValueError):
            confidence = -1.0
        if confidence >= 0:
            flags.append(f"ocr_confidence:{round(confidence, 3)}")
            if confidence < 0.6:
                flags.append("low_ocr_confidence")
    if len(re.sub(r"\s+", " ", (text or "").strip())) < 120:
        flags.append("short_chunk")
    if _count_sentence_marks(text) == 0:
        flags.append("no_sentence_punctuation")

    deduped: list[str] = []
    seen: set[str] = set()
    for flag in flags:
        if flag in seen:
            continue
        seen.add(flag)
        deduped.append(flag)
    return json.dumps(deduped, ensure_ascii=False)


def _chunk_semantic_units(
    units: list[str],
    *,
    document_id: str,
    source_path: str,
    source_title: str,
    page_number: int | None,
    starting_index: int,
    section_title: str | None = None,
    parent_section_title: str | None = None,
    extraction_method: str | None = None,
    has_text_layer: int | bool | None = None,
    ocr_confidence: float | None = None,
    created_at: str | None = None,
) -> tuple[list[dict], int]:
    target_chars = 1200
    overlap_units = 1
    created_at = created_at or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    chunks: list[dict] = []
    index = starting_index
    i = 0
    current_section = section_title
    current_parent_section = parent_section_title

    while i < len(units):
        current: list[str] = []
        char_count = 0
        local_section = current_section
        local_parent_section = current_parent_section
        start_i = i

        while i < len(units):
            unit = units[i].strip()
            if not unit:
                i += 1
                continue

            if _looks_like_heading(unit):
                heading = unit.rstrip(":").strip()
                if current:
                    # Heading encountered mid-chunk: flush the current chunk
                    # and let the next iteration handle the heading as the
                    # start of a new chunk.
                    break
                # Heading at the start of a new chunk: record it as section
                # context but do NOT append to `current`. Keeping the heading
                # out of chunk body text prevents heading+body concatenation
                # ("Cardholder Data Environment Defined The cardholder...")
                # from polluting embeddings and breaking PDF.js find()
                # highlighting downstream. The heading is preserved in the
                # section_title metadata field.
                local_section = heading or local_section
                local_parent_section = current_section if current_section and current_section != heading else current_parent_section
                i += 1
                continue

            if current and char_count + len(unit) > target_chars:
                break

            current.append(unit)
            char_count += len(unit) + 1
            i += 1

        chunk_text = " ".join(current).strip()
        if chunk_text:
            chunk_type = classify_chunk_type(
                chunk_text,
                section_title=local_section,
                page_number=page_number,
                extraction_method=extraction_method,
                has_text_layer=has_text_layer,
            )
            chunk_id = hashlib.sha256(
                f"{document_id}|{index}|{page_number or 'none'}|{source_path}|{chunk_text[:180]}".encode("utf-8")
            ).hexdigest()
            chunks.append({
                "chunk_id": chunk_id,
                "document_id": document_id,
                "chunk_index": index,
                "page_number": page_number,
                "section_title": local_section,
                "parent_section_title": local_parent_section,
                "chunk_type": chunk_type,
                "quality_flags": _chunk_quality_flags(
                    chunk_text,
                    chunk_type=chunk_type,
                    section_title=local_section,
                    page_number=page_number,
                    extraction_method=extraction_method,
                    has_text_layer=has_text_layer,
                    ocr_confidence=ocr_confidence,
                ),
                "text": chunk_text,
                "token_count": _token_count(chunk_text),
                "source_path": source_path,
                "created_at": created_at,
            })
            index += 1

        overlap = min(overlap_units, len(current) - 1) if len(current) >= 2 else 0
        i -= overlap
        if i <= start_i:
            i = start_i + 1
        current_section = local_section
        current_parent_section = local_parent_section

    return chunks, index


def _excel_col_letter(col_idx: int) -> str:
    letters = ""
    value = max(int(col_idx), 1)
    while value > 0:
        value, remainder = divmod(value - 1, 26)
        letters = chr(65 + remainder) + letters
    return letters


def _xlsx_safe_display_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    if hasattr(value, "isoformat"):
        try:
            return value.isoformat(sep=" ", timespec="seconds")
        except TypeError:
            return value.isoformat()
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.15g}"
    return str(value).strip()


def _xlsx_normalize_headers(values: list[Any], fallback_prefix: str = "Column") -> list[str]:
    headers: list[str] = []
    seen: dict[str, int] = {}
    for idx, raw_value in enumerate(values, start=1):
        label = _xlsx_safe_display_value(raw_value)
        if not label:
            label = f"{fallback_prefix} {_excel_col_letter(idx)}"
        label = re.sub(r"\s+", " ", label).strip()
        dedupe = label.lower()
        seen[dedupe] = seen.get(dedupe, 0) + 1
        if seen[dedupe] > 1:
            label = f"{label} {seen[dedupe]}"
        headers.append(label)
    return headers


def _xlsx_used_range(sheet) -> tuple[int, int, int, int] | None:
    try:
        dims = sheet.calculate_dimension()
        if not dims or ":" not in dims:
            return None
        from openpyxl.utils.cell import range_boundaries
        min_col, min_row, max_col, max_row = range_boundaries(dims)
        if min_col < 1 or min_row < 1:
            return None
        return min_row, max_row, min_col, max_col
    except Exception:
        return None


def _xlsx_row_values(row, min_col: int, max_col: int) -> list[Any]:
    values: list[Any] = []
    for col_idx in range(min_col, max_col + 1):
        cell = row[col_idx - min_col]
        values.append(cell.value if cell is not None else None)
    return values


def _xlsx_is_blank_row(values: list[Any]) -> bool:
    return all(_xlsx_safe_display_value(value) == "" for value in values)


def _xlsx_detect_header_row(rows: list[list[Any]]) -> int | None:
    if not rows:
        return None
    scan_limit = min(5, len(rows))
    for idx in range(scan_limit):
        row = rows[idx]
        if _xlsx_is_blank_row(row):
            continue
        text_cells = sum(1 for value in row if isinstance(value, str) and value.strip())
        non_empty = sum(1 for value in row if _xlsx_safe_display_value(value))
        numeric_cells = sum(1 for value in row if isinstance(value, (int, float)) and not isinstance(value, bool))
        if non_empty >= 2 and text_cells >= max(1, non_empty // 2) and text_cells >= numeric_cells:
            return idx
    return 0 if rows and not _xlsx_is_blank_row(rows[0]) else None


def _xlsx_row_text(headers: list[str], values: list[Any]) -> str:
    parts: list[str] = []
    for idx, value in enumerate(values):
        display = _xlsx_safe_display_value(value)
        if not display:
            continue
        label = headers[idx] if idx < len(headers) else f"Column {_excel_col_letter(idx + 1)}"
        parts.append(f"{label}: {display}")
    return " | ".join(parts)


def _xlsx_range_ref(min_row: int, max_row: int, min_col: int, max_col: int) -> str:
    return f"{_excel_col_letter(min_col)}{min_row}:{_excel_col_letter(max_col)}{max_row}"


def _xlsx_sheet_summary_text(workbook_name: str, sheet_name: str, range_ref: str, headers: list[str], max_row: int, max_col: int) -> str:
    header_text = ", ".join(headers[:8]) if headers else "No headers detected"
    return (
        f"Workbook: {workbook_name}\n"
        f"Sheet: {sheet_name}\n"
        f"Range: {range_ref}\n"
        f"Used range: {range_ref} ({max_row} rows x {max_col} columns)\n"
        f"Columns: {header_text}"
    )


def _workbook_structured_chunks(
    filepath: Path,
    *,
    document_id: str,
    source_title: str,
    file_type: str = "xlsx",
) -> tuple[list[dict], list[dict], int | None, str | None]:
    if load_workbook is None:
        return [], [], None, "Spreadsheet support requires openpyxl"

    try:
        workbook = load_workbook(filepath, read_only=True, data_only=True)
    except Exception as exc:
        return [], [], None, f"Failed to read {file_type.upper()}: {exc}"

    source_path = _relative_source_path(filepath)
    workbook_name = filepath.name
    chunks: list[dict] = []
    sheets: list[dict] = []
    chunk_index = 0
    total_rows = 0

    try:
        for sheet in workbook.worksheets:
            if getattr(sheet, "sheet_state", "visible") != "visible":
                continue

            used = _xlsx_used_range(sheet)
            if not used:
                sheets.append({
                    "sheet_name": sheet.title,
                    "used_range": None,
                    "max_row": int(getattr(sheet, "max_row", 0) or 0),
                    "max_column": int(getattr(sheet, "max_column", 0) or 0),
                    "headers": [],
                    "rows": [],
                })
                continue

            min_row, max_row, min_col, max_col = used
            raw_rows: list[list[Any]] = []
            for row in sheet.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                values = _xlsx_row_values(row, min_col, max_col)
                if _xlsx_is_blank_row(values):
                    continue
                raw_rows.append(values)

            if not raw_rows:
                sheets.append({
                    "sheet_name": sheet.title,
                    "used_range": _xlsx_range_ref(min_row, max_row, min_col, max_col),
                    "max_row": int(getattr(sheet, "max_row", 0) or 0),
                    "max_column": int(getattr(sheet, "max_column", 0) or 0),
                    "headers": [],
                    "rows": [],
                })
                continue

            header_idx = _xlsx_detect_header_row(raw_rows)
            if header_idx is not None:
                headers = _xlsx_normalize_headers(raw_rows[header_idx], fallback_prefix="Column")
                data_rows = raw_rows[header_idx + 1:]
                header_row_number = min_row + header_idx
            else:
                headers = [f"Column {_excel_col_letter(idx)}" for idx in range(min_col, max_col + 1)]
                data_rows = raw_rows
                header_row_number = None

            sheet_range_ref = _xlsx_range_ref(min_row, max_row, min_col, max_col)
            summary_text = _xlsx_sheet_summary_text(workbook_name, sheet.title, sheet_range_ref, headers, int(getattr(sheet, "max_row", 0) or max_row), int(getattr(sheet, "max_column", 0) or max_col))
            sheet_rows: list[dict] = []
            for row_offset, row_values in enumerate(data_rows, start=1):
                row_number = (min_row + header_idx + row_offset) if header_idx is not None else (min_row + row_offset - 1)
                row_text = _xlsx_row_text(headers, row_values)
                if not row_text:
                    continue
                sheet_rows.append({
                    "row_number": row_number,
                    "text": row_text,
                })

            sheets.append({
                "sheet_name": sheet.title,
                "used_range": sheet_range_ref,
                "max_row": int(getattr(sheet, "max_row", 0) or max_row),
                "max_column": int(getattr(sheet, "max_column", 0) or max_col),
                "headers": headers,
                "rows": sheet_rows,
            })
            total_rows += len(sheet_rows)

            summary_chunk_id = hashlib.sha256(f"{document_id}|{sheet.title}|summary|{sheet_range_ref}".encode("utf-8")).hexdigest()
            chunks.append({
                "chunk_id": summary_chunk_id,
                "document_id": document_id,
                "chunk_index": chunk_index,
                "page_number": None,
                "section_title": sheet.title,
                "parent_section_title": workbook_name,
                "chunk_type": "spreadsheet_sheet_summary",
                "quality_flags": json.dumps(["spreadsheet", "sheet_summary"], ensure_ascii=False),
                "extraction_method": "xlsx_openpyxl" if file_type == "xlsx" else "xlsm_openpyxl",
                "has_text_layer": 1,
                "ocr_confidence": None,
                "char_count": len(summary_text),
                "word_count": _token_count(summary_text),
                "text": summary_text,
                "token_count": _token_count(summary_text),
                "source_path": source_path,
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "sheet_name": sheet.title,
                "range_ref": sheet_range_ref,
                "row_start": min_row,
                "row_end": max_row,
                "col_start": min_col,
                "col_end": max_col,
                "headers": headers,
                "workbook_name": workbook_name,
                "has_structured_preview": 1,
            })
            chunk_index += 1

            table_text = (
                f"Workbook: {workbook_name}\n"
                f"Sheet: {sheet.title}\n"
                f"Range: {sheet_range_ref}\n"
                f"Columns: {', '.join(headers[:12])}"
            )
            if len(sheet_rows) > 0:
                table_text += f"\nRows: {sheet_rows[0]['row_number']}-{sheet_rows[-1]['row_number']}"
            table_chunk_id = hashlib.sha256(f"{document_id}|{sheet.title}|table|{sheet_range_ref}".encode("utf-8")).hexdigest()
            chunks.append({
                "chunk_id": table_chunk_id,
                "document_id": document_id,
                "chunk_index": chunk_index,
                "page_number": None,
                "section_title": sheet.title,
                "parent_section_title": workbook_name,
                "chunk_type": "spreadsheet_table_summary",
                "quality_flags": json.dumps(["spreadsheet", "table_summary"], ensure_ascii=False),
                "extraction_method": "xlsx_openpyxl" if file_type == "xlsx" else "xlsm_openpyxl",
                "has_text_layer": 1,
                "ocr_confidence": None,
                "char_count": len(table_text),
                "word_count": _token_count(table_text),
                "text": table_text,
                "token_count": _token_count(table_text),
                "source_path": source_path,
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "sheet_name": sheet.title,
                "range_ref": sheet_range_ref,
                "row_start": min_row,
                "row_end": max_row,
                "col_start": min_col,
                "col_end": max_col,
                "headers": headers,
                "workbook_name": workbook_name,
                "has_structured_preview": 1,
            })
            chunk_index += 1

            batch_size = 8
            for batch_start in range(0, len(sheet_rows), batch_size):
                batch = sheet_rows[batch_start:batch_start + batch_size]
                if not batch:
                    continue
                row_start = batch[0]["row_number"]
                row_end = batch[-1]["row_number"]
                batch_range_ref = _xlsx_range_ref(row_start, row_end, min_col, max_col)
                row_lines = [f"- Row {item['row_number']}: {item['text']}" for item in batch]
                batch_text = (
                    f"Workbook: {workbook_name}\n"
                    f"Sheet: {sheet.title}\n"
                    f"Range: {batch_range_ref}\n"
                    f"Columns: {', '.join(headers)}\n"
                    f"Rows {row_start}-{row_end}:\n" + "\n".join(row_lines)
                )
                batch_chunk_id = hashlib.sha256(f"{document_id}|{sheet.title}|row_batch|{batch_range_ref}|{batch_start}".encode("utf-8")).hexdigest()
                chunks.append({
                    "chunk_id": batch_chunk_id,
                    "document_id": document_id,
                    "chunk_index": chunk_index,
                    "page_number": None,
                    "section_title": sheet.title,
                    "parent_section_title": workbook_name,
                    "chunk_type": "spreadsheet_row_batch",
                    "quality_flags": json.dumps(["spreadsheet", "row_batch"], ensure_ascii=False),
                    "extraction_method": "xlsx_openpyxl" if file_type == "xlsx" else "xlsm_openpyxl",
                    "has_text_layer": 1,
                    "ocr_confidence": None,
                    "char_count": len(batch_text),
                    "word_count": _token_count(batch_text),
                    "text": batch_text,
                    "token_count": _token_count(batch_text),
                    "source_path": source_path,
                    "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "sheet_name": sheet.title,
                    "range_ref": batch_range_ref,
                    "row_start": row_start,
                    "row_end": row_end,
                    "col_start": min_col,
                    "col_end": max_col,
                    "headers": headers,
                    "workbook_name": workbook_name,
                    "has_structured_preview": 1,
                })
                chunk_index += 1

            for row_item in sheet_rows:
                if len(row_item["text"]) < 40:
                    continue
                kv_text = (
                    f"Workbook: {workbook_name}\n"
                    f"Sheet: {sheet.title}\n"
                    f"Range: {_xlsx_range_ref(row_item['row_number'], row_item['row_number'], min_col, max_col)}\n"
                    f"Row {row_item['row_number']}:\n{row_item['text']}"
                )
                kv_chunk_id = hashlib.sha256(f"{document_id}|{sheet.title}|row_kv|{row_item['row_number']}".encode("utf-8")).hexdigest()
                chunks.append({
                    "chunk_id": kv_chunk_id,
                    "document_id": document_id,
                    "chunk_index": chunk_index,
                    "page_number": None,
                    "section_title": sheet.title,
                    "parent_section_title": workbook_name,
                    "chunk_type": "spreadsheet_key_value",
                    "quality_flags": json.dumps(["spreadsheet", "key_value"], ensure_ascii=False),
                    "extraction_method": "xlsx_openpyxl" if file_type == "xlsx" else "xlsm_openpyxl",
                    "has_text_layer": 1,
                    "ocr_confidence": None,
                    "char_count": len(kv_text),
                    "word_count": _token_count(kv_text),
                    "text": kv_text,
                    "token_count": _token_count(kv_text),
                    "source_path": source_path,
                    "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "sheet_name": sheet.title,
                    "range_ref": _xlsx_range_ref(row_item["row_number"], row_item["row_number"], min_col, max_col),
                    "row_start": row_item["row_number"],
                    "row_end": row_item["row_number"],
                    "col_start": min_col,
                    "col_end": max_col,
                    "headers": headers,
                    "workbook_name": workbook_name,
                    "has_structured_preview": 1,
                })
                chunk_index += 1
    finally:
        workbook.close()

    if not chunks:
        return [], sheets, None, f"No extractable {file_type.upper()} content found"
    return chunks, sheets, total_rows, None


def _xlsx_structured_chunks(filepath: Path, *, document_id: str, source_title: str) -> tuple[list[dict], list[dict], int | None, str | None]:
    return _workbook_structured_chunks(
        filepath,
        document_id=document_id,
        source_title=source_title,
        file_type="xlsx",
    )


def _xlsm_structured_chunks(filepath: Path, *, document_id: str, source_title: str) -> tuple[list[dict], list[dict], int | None, str | None]:
    return _workbook_structured_chunks(
        filepath,
        document_id=document_id,
        source_title=source_title,
        file_type="xlsm",
    )


def _detect_csv_dialect(sample: str) -> csv.Dialect:
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        return dialect
    except Exception:
        class _FallbackDialect(csv.Dialect):
            delimiter = ","
            quotechar = '"'
            doublequote = True
            skipinitialspace = True
            lineterminator = "\n"
            quoting = csv.QUOTE_MINIMAL
        return _FallbackDialect()


def _csv_structured_chunks(filepath: Path, *, document_id: str, source_title: str) -> tuple[list[dict], list[dict], int | None, str | None]:
    source_path = _relative_source_path(filepath)
    workbook_name = filepath.name
    sheet_name = filepath.stem or "Sheet1"
    try:
        raw = filepath.read_text(encoding="utf-8-sig", errors="ignore")
    except Exception as exc:
        return [], [], None, f"Failed to read CSV: {exc}"

    if not raw.strip():
        return [], [], None, "No extractable CSV content found"

    lines = raw.splitlines()
    sample = "\n".join(lines[:20]) if lines else raw[:4096]
    dialect = _detect_csv_dialect(sample)

    try:
        reader = csv.reader(lines, dialect=dialect)
        raw_rows = [[str(cell).strip() for cell in row] for row in reader]
    except Exception as exc:
        return [], [], None, f"Failed to parse CSV: {exc}"

    raw_rows = [row for row in raw_rows if any((cell or "").strip() for cell in row)]
    if not raw_rows:
        return [], [], None, "No extractable CSV rows found"

    max_col = max((len(row) for row in raw_rows), default=0)
    if max_col <= 0:
        return [], [], None, "No extractable CSV columns found"

    padded_rows = [row + [""] * (max_col - len(row)) for row in raw_rows]
    header_idx = _xlsx_detect_header_row(padded_rows)
    if header_idx is not None:
        headers = _xlsx_normalize_headers(padded_rows[header_idx], fallback_prefix="Column")
        data_rows = padded_rows[header_idx + 1:]
        first_data_row_number = header_idx + 2
    else:
        headers = [f"Column {_excel_col_letter(idx)}" for idx in range(1, max_col + 1)]
        data_rows = padded_rows
        first_data_row_number = 1

    sheet_rows: list[dict] = []
    for offset, values in enumerate(data_rows):
        row_number = first_data_row_number + offset
        row_text = _xlsx_row_text(headers, values)
        if not row_text:
            continue
        sheet_rows.append({"row_number": row_number, "text": row_text, "values": values})

    min_col = 1
    max_row = len(padded_rows)
    sheet_range_ref = _xlsx_range_ref(1, max_row, min_col, max_col)
    chunks: list[dict] = []
    chunk_index = 0

    summary_text = _xlsx_sheet_summary_text(workbook_name, sheet_name, sheet_range_ref, headers, max_row, max_col)
    summary_chunk_id = hashlib.sha256(f"{document_id}|{sheet_name}|summary|{sheet_range_ref}".encode("utf-8")).hexdigest()
    chunks.append({
        "chunk_id": summary_chunk_id,
        "document_id": document_id,
        "chunk_index": chunk_index,
        "page_number": None,
        "section_title": sheet_name,
        "parent_section_title": workbook_name,
        "chunk_type": "spreadsheet_sheet_summary",
        "quality_flags": json.dumps(["spreadsheet", "sheet_summary"], ensure_ascii=False),
        "extraction_method": "csv_parser",
        "has_text_layer": 1,
        "ocr_confidence": None,
        "char_count": len(summary_text),
        "word_count": _token_count(summary_text),
        "text": summary_text,
        "token_count": _token_count(summary_text),
        "source_path": source_path,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "sheet_name": sheet_name,
        "range_ref": sheet_range_ref,
        "row_start": 1,
        "row_end": max_row,
        "col_start": min_col,
        "col_end": max_col,
        "headers": headers,
        "workbook_name": workbook_name,
        "has_structured_preview": 1,
    })
    chunk_index += 1

    table_text = (
        f"Workbook: {workbook_name}\n"
        f"Sheet: {sheet_name}\n"
        f"Range: {sheet_range_ref}\n"
        f"Columns: {', '.join(headers[:12])}"
    )
    if sheet_rows:
        table_text += f"\nRows: {sheet_rows[0]['row_number']}-{sheet_rows[-1]['row_number']}"
    table_chunk_id = hashlib.sha256(f"{document_id}|{sheet_name}|table|{sheet_range_ref}".encode("utf-8")).hexdigest()
    chunks.append({
        "chunk_id": table_chunk_id,
        "document_id": document_id,
        "chunk_index": chunk_index,
        "page_number": None,
        "section_title": sheet_name,
        "parent_section_title": workbook_name,
        "chunk_type": "spreadsheet_table_summary",
        "quality_flags": json.dumps(["spreadsheet", "table_summary"], ensure_ascii=False),
        "extraction_method": "csv_parser",
        "has_text_layer": 1,
        "ocr_confidence": None,
        "char_count": len(table_text),
        "word_count": _token_count(table_text),
        "text": table_text,
        "token_count": _token_count(table_text),
        "source_path": source_path,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "sheet_name": sheet_name,
        "range_ref": sheet_range_ref,
        "row_start": 1,
        "row_end": max_row,
        "col_start": min_col,
        "col_end": max_col,
        "headers": headers,
        "workbook_name": workbook_name,
        "has_structured_preview": 1,
    })
    chunk_index += 1

    batch_size = 12
    for batch_start in range(0, len(sheet_rows), batch_size):
        batch = sheet_rows[batch_start:batch_start + batch_size]
        if not batch:
            continue
        row_start = batch[0]["row_number"]
        row_end = batch[-1]["row_number"]
        batch_range_ref = _xlsx_range_ref(row_start, row_end, min_col, max_col)
        row_lines = [f"- Row {item['row_number']}: {item['text']}" for item in batch]
        batch_text = (
            f"Workbook: {workbook_name}\n"
            f"Sheet: {sheet_name}\n"
            f"Range: {batch_range_ref}\n"
            f"Columns: {', '.join(headers)}\n"
            f"Rows {row_start}-{row_end}:\n" + "\n".join(row_lines)
        )
        batch_chunk_id = hashlib.sha256(f"{document_id}|{sheet_name}|row_batch|{batch_range_ref}|{batch_start}".encode("utf-8")).hexdigest()
        chunks.append({
            "chunk_id": batch_chunk_id,
            "document_id": document_id,
            "chunk_index": chunk_index,
            "page_number": None,
            "section_title": sheet_name,
            "parent_section_title": workbook_name,
            "chunk_type": "spreadsheet_row_batch",
            "quality_flags": json.dumps(["spreadsheet", "row_batch"], ensure_ascii=False),
            "extraction_method": "csv_parser",
            "has_text_layer": 1,
            "ocr_confidence": None,
            "char_count": len(batch_text),
            "word_count": _token_count(batch_text),
            "text": batch_text,
            "token_count": _token_count(batch_text),
            "source_path": source_path,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "sheet_name": sheet_name,
            "range_ref": batch_range_ref,
            "row_start": row_start,
            "row_end": row_end,
            "col_start": min_col,
            "col_end": max_col,
            "headers": headers,
            "workbook_name": workbook_name,
            "has_structured_preview": 1,
        })
        chunk_index += 1

    sheets = [{
        "sheet_name": sheet_name,
        "used_range": sheet_range_ref,
        "max_row": max_row,
        "max_column": max_col,
        "headers": headers,
        "rows": [{"row_number": row["row_number"], "text": row["text"]} for row in sheet_rows],
    }]
    return chunks, sheets, len(sheet_rows), None


def _extract_pdf_pages(filepath: Path) -> list[dict[str, Any]]:
    pages = _extract_pdf_page_records(filepath)
    for page in pages:
        text = page.get("text", "")
        units = split_semantic_units(text)
        if not units and text.strip():
            units = [text.strip()]
        page["units"] = units
        page["section_title"] = units[0] if units and _looks_like_heading(units[0]) else None
    return pages


def _extract_docx_segments(filepath: Path) -> list[dict]:
    doc = DocxDocument(filepath)
    units = [line.strip() for line in _iter_docx_text(doc) if line.strip()]
    if not units:
        return []
    return [{
        "page_number": None,
        "text": "\n".join(units),
        "units": units,
        "section_title": units[0] if _looks_like_heading(units[0]) else None,
    }]


def _extract_xlsx_segments(filepath: Path) -> tuple[list[dict], str | None]:
    try:
        from openpyxl import load_workbook
    except Exception:
        return [], "XLSX support requires openpyxl"

    try:
        workbook = load_workbook(filepath, read_only=True, data_only=True)
    except Exception as exc:
        return [], f"Failed to read XLSX: {exc}"

    segments: list[dict] = []
    for sheet in workbook.worksheets:
        if getattr(sheet, "sheet_state", "visible") != "visible":
            continue
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell not in {None, ""}]
            if values:
                rows.append(" | ".join(values))
        if rows:
            segments.append({
                "page_number": None,
                "text": "\n".join(rows),
                "units": rows,
                "section_title": sheet.title,
            })
    return segments, None


def _build_chunks_for_file(filepath: Path, *, document_id: str, source_title: str) -> tuple[list[dict], list[dict], int | None, str | None]:
    ext = filepath.suffix.lower()
    source_path = _relative_source_path(filepath)

    if ext == ".pdf":
        pages = _extract_pdf_pages(filepath)
        chunks: list[dict] = []
        next_index = 0
        for page in pages:
            if not page.get("text", "").strip():
                continue
            segment_chunks, next_index = _chunk_semantic_units(
                page["units"],
                document_id=document_id,
                source_path=source_path,
                source_title=source_title,
                page_number=page["page_number"],
                starting_index=next_index,
                section_title=page.get("section_title"),
                extraction_method=page.get("extraction_method"),
                has_text_layer=page.get("has_text_layer", 1),
                ocr_confidence=page.get("ocr_confidence"),
            )
            for chunk in segment_chunks:
                chunk["extraction_method"] = page.get("extraction_method")
                chunk["has_text_layer"] = page.get("has_text_layer", 1)
                chunk["ocr_confidence"] = page.get("ocr_confidence")
                chunk["char_count"] = page.get("char_count")
                chunk["word_count"] = page.get("word_count")
            chunks.extend(segment_chunks)
        return chunks, pages, len(pages), None

    if ext == ".docx":
        segments = _extract_docx_segments(filepath)
        if not segments:
            return [], [], None, "No extractable DOCX text found"
        chunks: list[dict] = []
        next_index = 0
        for segment in segments:
            segment_chunks, next_index = _chunk_semantic_units(
                segment["units"],
                document_id=document_id,
                source_path=source_path,
                source_title=source_title,
                page_number=None,
                starting_index=next_index,
                section_title=segment.get("section_title"),
            )
            chunks.extend(segment_chunks)
        return chunks, [], None, None

    if ext == ".xlsx":
        return _xlsx_structured_chunks(filepath, document_id=document_id, source_title=source_title)
    if ext == ".xlsm":
        return _xlsm_structured_chunks(filepath, document_id=document_id, source_title=source_title)
    if ext == ".csv":
        return _csv_structured_chunks(filepath, document_id=document_id, source_title=source_title)
    if ext == ".xls":
        return [], [], None, "Unsupported .xls format. Convert to .xlsx or .csv and re-ingest."
    if ext == ".ods":
        return [], [], None, "Unsupported .ods format. Convert to .xlsx or .csv and re-ingest."
    units = split_semantic_units(extract_text(filepath))
    if not units:
        return [], [], None, "No extractable text found"
    chunks, _ = _chunk_semantic_units(
        units,
        document_id=document_id,
        source_path=source_path,
        source_title=source_title,
        page_number=None,
        starting_index=0,
        extraction_method=None,
        has_text_layer=None,
        ocr_confidence=None,
    )
    return chunks, [], None, None


def _delete_chroma_source(collection, source_path: str) -> None:
    try:
        existing = collection.get(where={"source_path": source_path})
        if existing and existing.get("ids"):
            collection.delete(ids=existing["ids"])
    except Exception as exc:
        print(f"[Spark Ingest]   Failed to delete Chroma records for {source_path}: {exc}")


def _sanitize_chroma_metadata(metadata: dict[str, Any]) -> dict[str, Any]:
    clean: dict[str, Any] = {}
    for key, value in metadata.items():
        if value is None:
            continue
        if isinstance(value, bool):
            clean[key] = int(value)
            continue
        if isinstance(value, (str, int, float)):
            clean[key] = value
            continue
        clean[key] = str(value)
    return clean


DEPARTMENT_MAP = {
    "HR":         ["PTO", "Benefits", "Employee", "Leave", "Handbook",
                   "Onboarding", "Performance", "Compensation"],
    "IT":         ["IT", "Support", "Security", "Technology", "Data",
                   "Governance", "Enterprise"],
    "Finance":    ["Finance", "Budget", "Expense", "Accounting", "Invoice"],
    "Compliance": ["Compliance", "Audit", "Risk", "Regulatory", "Policy"],
    "Operations": ["Operations", "Facilities", "Branch", "Procedure"],
    "Credit":     ["Credit", "Loan", "Mortgage", "PHH", "Underwriting"],
    "Marketing":  ["Marketing", "Brand", "Communications"],
    "Legal":      ["Legal", "Contract", "Agreement"],
}

def _infer_department(filename: str) -> str:
    name_upper = filename.upper()
    for dept, keywords in DEPARTMENT_MAP.items():
        for kw in keywords:
            if kw.upper() in name_upper:
                return dept
    return "General"


def _get_chroma_collection():
    client = chromadb.PersistentClient(path=str(CHROMA_FOLDER))
    return client.get_or_create_collection(
        name=CHROMA_COLLECTION,
        metadata={"hnsw:space": "cosine"}
    )


def load_hashes() -> dict:
    if HASH_STORE.exists():
        with open(HASH_STORE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_hashes(hashes: dict):
    CHROMA_FOLDER.mkdir(parents=True, exist_ok=True)
    with open(HASH_STORE, "w", encoding="utf-8") as f:
        json.dump(hashes, f, indent=2)

def file_hash(filepath: Path) -> str:
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def clean_text(text: str) -> str:
    text = re.sub(r"^[\s=\-~#*]{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 90:
        return False
    if stripped.endswith(":"):
        return True
    words = stripped.split()
    if len(words) <= 10 and stripped.upper() == stripped:
        return True
    alpha = [w for w in words if re.search(r"[A-Za-z]", w)]
    if 0 < len(alpha) <= 8 and all(w[:1].isupper() for w in alpha):
        return True
    return False


def _is_structured_line(line: str) -> bool:
    stripped = line.strip()
    return bool(
        re.match(r"^(?:[-*•]|(\d+[\.)]))\s+", stripped)
        or "|" in stripped
        or re.match(r"^[A-Z][A-Z0-9 _/-]{3,}$", stripped)
    )


def _read_text_file(filepath: Path) -> str:
    return filepath.read_text(encoding="utf-8", errors="ignore")


def _iter_docx_text(doc: DocxDocument) -> Iterable[str]:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                yield " | ".join(cells)


def _read_docx_file(filepath: Path) -> str:
    doc = DocxDocument(filepath)
    return "\n".join(_iter_docx_text(doc))


def _normalize_pdf_page(text: str) -> str:
    lines = text.split("\n")
    rebuilt = []

    for line in lines:
        stripped = line.rstrip()
        if not stripped:
            rebuilt.append("")
            continue

        if rebuilt and not re.search(r"[.!?:]\s*$", rebuilt[-1]) and rebuilt[-1] != "":
            rebuilt[-1] = rebuilt[-1] + " " + stripped
        else:
            rebuilt.append(stripped)

    text_out = "\n".join(rebuilt)
    text_out = re.sub(r"\n{3,}", "\n\n", text_out)
    text_out = re.sub(r" {2,}", " ", text_out)
    return text_out.strip()


def _count_meaningful_chars(text: str) -> int:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    return len(re.sub(r"[^A-Za-z0-9]", "", cleaned))


def _count_meaningful_words(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text or ""))


def _page_needs_ocr(text: str) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    return _count_meaningful_chars(cleaned) < PDF_OCR_CHAR_THRESHOLD or _count_meaningful_words(cleaned) < PDF_OCR_WORD_THRESHOLD


def _read_pdf_file(filepath: Path) -> str:
    reader = PdfReader(str(filepath))
    pages = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        page_text = _normalize_pdf_page(page_text)
        if page_text:
            pages.append(page_text)

    return "\n\n".join(pages)


def _extract_pdf_page_records(filepath: Path) -> list[dict[str, Any]]:
    reader = PdfReader(str(filepath))
    page_records: list[dict[str, Any]] = []

    for page_number, page in enumerate(reader.pages, start=1):
        raw_text = _normalize_pdf_page(page.extract_text() or "")
        has_text_layer = bool(raw_text and not _page_needs_ocr(raw_text))
        if has_text_layer:
            normalized_text = raw_text
            method = "text_layer"
            confidence = None
        else:
            print(f"[Spark OCR] Page {page_number} has weak text layer; running Tesseract OCR")
            ocr_result = extract_pdf_page_text(filepath, page_number)
            normalized_text = _normalize_pdf_page(ocr_result.get("text", "") or "")
            confidence = ocr_result.get("confidence")
            method = str(ocr_result.get("method") or "ocr")
            if method not in {"ocr", "ocr_failed", "vision_fallback"}:
                method = "ocr"
            if method == "ocr":
                ocr_chars = _count_meaningful_chars(normalized_text)
                ocr_words = _count_meaningful_words(normalized_text)
                print(f"[Spark OCR] Page {page_number} OCR extracted {ocr_chars} chars / {ocr_words} words")

        char_count = _count_meaningful_chars(normalized_text)
        word_count = _count_meaningful_words(normalized_text)
        page_records.append({
            "page_number": page_number,
            "text": normalized_text,
            "extraction_method": method,
            "has_text_layer": int(has_text_layer),
            "ocr_confidence": confidence,
            "char_count": char_count,
            "word_count": word_count,
        })

    return page_records


def _read_rtf_file(filepath: Path) -> str:
    raw = filepath.read_text(encoding="utf-8", errors="ignore")
    raw = re.sub(r"\\par[d]?", "\n", raw)
    raw = re.sub(r"\\tab", "\t", raw)
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", "", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d*\s?", "", raw)
    raw = raw.replace("{", "").replace("}", "")
    return raw


def extract_text(filepath: Path) -> str:
    ext = filepath.suffix.lower()

    try:
        if ext in {".txt", ".md"}:
            raw = _read_text_file(filepath)
        elif ext == ".docx":
            raw = _read_docx_file(filepath)
        elif ext == ".pdf":
            raw = _read_pdf_file(filepath)
        elif ext == ".rtf":
            raw = _read_rtf_file(filepath)
        else:
            return ""
    except Exception as e:
        print(f"[Spark Ingest]   Failed to extract text from {filepath.name}: {e}")
        return ""

    return clean_text(raw)


def _hash_key(filepath: Path) -> str:
    try:
        return filepath.relative_to(INTAKE_FOLDER).as_posix()
    except ValueError:
        return filepath.name


def split_sentences(text: str) -> list:
    paragraphs = re.split(r"\n{2,}", text)
    sentences  = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        parts = re.split(r"(?<=[a-zA-Z][.!?])\s+", para)
        for part in parts:
            part = part.strip()
            if part and len(part) > 10:
                sentences.append(part)
    return sentences


def split_semantic_units(text: str) -> list[str]:
    paragraphs = re.split(r"\n{2,}", text)
    units: list[str] = []

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        lines = [line.strip() for line in para.splitlines() if line.strip()]
        if not lines:
            continue

        if len(lines) == 1:
            line = lines[0]
            if _looks_like_heading(line) or _is_structured_line(line):
                units.append(line)
            else:
                units.extend(split_sentences(line))
            continue

        if any(_is_structured_line(line) for line in lines):
            units.append(" ".join(lines))
            continue

        if _looks_like_heading(lines[0]):
            units.append(lines[0])
            remainder = " ".join(lines[1:]).strip()
            if remainder:
                units.extend(split_sentences(remainder))
            continue

        units.extend(split_sentences(" ".join(lines)))

    return [unit.strip() for unit in units if unit.strip()]

def chunk_text(text: str, source: str) -> list:
    TARGET_CHARS  = 1200
    OVERLAP_UNITS = 2

    units = split_semantic_units(text)
    if not units:
        return []

    chunks = []
    idx    = 0
    i      = 0

    while i < len(units):
        current:    list[str] = []
        char_count: int       = 0
        i_start = i                     

        while i < len(units):
            s = units[i]
            if char_count + len(s) > TARGET_CHARS and current:
                break
            current.append(s)
            char_count += len(s) + 1
            i += 1

        chunk_val = " ".join(current).strip()
        if chunk_val:
            chunk_id = hashlib.sha256(f"{source}_{idx}_{chunk_val[:128]}".encode()).hexdigest()
            chunks.append({
                "id":          chunk_id,
                "source":      source,
                "chunk_index": idx,
                "text":        chunk_val,
            })
            idx += 1

        effective_overlap = min(OVERLAP_UNITS, len(current) - 1) if len(current) >= OVERLAP_UNITS else 0
        i -= effective_overlap
        if i <= i_start:
            i = i_start + 1

    return chunks


def _build_source_fingerprint(filepath: Path, current_hash: str) -> str:
    stat = filepath.stat()
    payload = f"{filepath.name}|{stat.st_size}|{int(stat.st_mtime)}|{current_hash}"
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()

# ── Main ─────────────────────────────────────────────────────────

def run_ingestion(target_source_path: str | None = None, *, force_reingest: bool = False):
    init_db()
    print(f"[Spark Ingest] Initializing ChromaDB at: {CHROMA_FOLDER} collection={CHROMA_COLLECTION}")
    collection = _get_chroma_collection()
    active_embedding_config = get_active_embedding_config()
    print(
        f"[Spark Ingest] Embedding API configured: "
        f"{active_embedding_config.get('embedding_model_id')} "
        f"provider={active_embedding_config.get('embedding_provider')} "
        f"base={active_embedding_config.get('embedding_base_url') or active_embedding_config.get('embedding_model_path')} "
        f"normalized={active_embedding_config.get('embedding_normalized')}"
    )
    manifest = _load_manifest_metadata()

    files = [
        f for f in INTAKE_FOLDER.rglob("*")
        if f.is_file() and (f.suffix.lower() in SUPPORTED_EXTENSIONS or f.suffix.lower() in DISABLED_EXTENSIONS)
    ]
    if target_source_path:
        target_norm = str(target_source_path).replace("\\", "/").lstrip("./")
        target_norm = target_norm.lower()
        files = [
            f for f in files
            if _relative_source_path(f).replace("\\", "/").lower() == target_norm
            or f.name.lower() == target_norm
        ]
    current_source_paths = {_relative_source_path(f) for f in files}
    hashes = load_hashes()
    run_id = start_ingestion_run()

    total_chunks = 0
    skipped = 0
    ingested = 0
    failed = 0
    processed = 0
    failed_items: list[dict[str, str]] = []
    skipped_items: list[str] = []
    unsupported_items: list[dict[str, str]] = []

    try:
        if not files:
            print("[Spark Ingest] Nothing to ingest.")
        
        seen_text_counts: dict[str, int] = {}
        total_vectors = 0
        vector_skipped_run = 0

        for filepath in files:
            processed += 1
            source_path = _relative_source_path(filepath)
            current_hash = file_hash(filepath)
            existing = get_document_by_source_path(source_path)
            existing_chunk_count = get_document_chunk_count(existing["document_id"]) if existing else 0

            if (
                not force_reingest
                and existing
                and str(existing.get("file_hash")) == current_hash
                and str(existing.get("status", "active")) == "active"
                and existing_chunk_count > 0
            ):
                print(f"[Spark Ingest]   Skipping (unchanged): {filepath.name}")
                skipped += 1
                skipped_items.append(source_path)
                hashes[source_path] = current_hash
                continue

            if existing and existing_chunk_count == 0:
                print(f"[Spark Ingest]   Reprocessing zero-chunk document: {filepath.name}")

            print(f"[Spark Ingest] Processing: {filepath.name}")
            if existing:
                delete_document_chunks(existing["document_id"])
                delete_document_pages(existing["document_id"])
                _delete_chroma_source(collection, source_path)

            manifest_row = _lookup_manifest_metadata(filepath, manifest)
            document = _normalize_manifest_row(filepath, manifest_row, current_hash)
            document["document_id"] = existing["document_id"] if existing else document["document_id"]
            document["last_ingested_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            if filepath.suffix.lower() in DISABLED_EXTENSIONS:
                document["status"] = "inactive"
                document["ingestion_status"] = "unsupported"
                document["ingestion_error"] = f"Ingestion for {filepath.suffix.lower()} is currently disabled."
                upsert_document(document)
                failed += 1
                unsupported_items.append({"source_path": source_path, "reason": document["ingestion_error"]})
                print(f"[Spark Ingest]   {document['ingestion_error']}: {filepath.name}")
                continue

            if filepath.suffix.lower() not in SUPPORTED_EXTENSIONS:
                document["status"] = "inactive"
                document["ingestion_status"] = "unsupported"
                document["ingestion_error"] = "Unsupported file type"
                upsert_document(document)
                failed += 1
                unsupported_items.append({"source_path": source_path, "reason": document["ingestion_error"]})
                continue

            try:
                chunks, pages, page_count, error = _build_chunks_for_file(
                    filepath,
                    document_id=document["document_id"],
                    source_title=document["title"],
                )
            except Exception as exc:
                chunks, pages, page_count, error = [], [], None, str(exc)

            document["page_count"] = page_count
            document["status"] = "active"
            document["ingestion_status"] = "ingesting"
            document["ingestion_error"] = None
            document = upsert_document(document)

            if pages:
                insert_document_pages(document["document_id"], pages)

            if error or not chunks:
                document["status"] = "inactive"
                normalized_error = (error or "").lower()
                document["ingestion_status"] = "unsupported" if "unsupported" in normalized_error else "failed"
                document["ingestion_error"] = error or "No chunks created"
                upsert_document(document)
                failed += 1
                failed_items.append({"source_path": source_path, "reason": document["ingestion_error"]})
                print(f"[Spark Ingest]   {document['ingestion_error']}: {filepath.name}")
                continue

            # Classify vector eligibility for each chunk
            vector_eligible_chunks = []
            file_embedding_created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            for c in chunks:
                is_eligible, skip_reason = classify_vector_eligibility(c, seen_text_counts)
                c["vector_eligible"] = 1 if is_eligible else 0
                c["vector_skip_reason"] = skip_reason
                if is_eligible:
                    c["embedding_model_id"] = active_embedding_config["embedding_model_id"]
                    c["embedding_provider"] = active_embedding_config["embedding_provider"]
                    c["embedding_base_url"] = active_embedding_config["embedding_base_url"]
                    c["embedding_model_path"] = active_embedding_config.get("embedding_model_path") or ""
                    c["embedding_normalized"] = active_embedding_config["embedding_normalized"]
                    c["embedding_instruction"] = active_embedding_config["embedding_instruction"]
                    c["embedding_config_hash"] = active_embedding_config["embedding_config_hash"]
                    c["embedding_created_at"] = file_embedding_created_at
                    vector_eligible_chunks.append(c)
                else:
                    vector_skipped_run += 1

            embeds = []
            if vector_eligible_chunks:
                embed_texts = [f"[Document: {document['title']}] {c['text']}" for c in vector_eligible_chunks]
                embeds = call_embedding_api(embed_texts, is_query=False)
                embedding_dimension = len(embeds[0]) if len(embeds) > 0 else 0
                for c in vector_eligible_chunks:
                    c["embedding_dimension"] = embedding_dimension

                _delete_chroma_source(collection, source_path)
                metadatas = []
                for c in vector_eligible_chunks:
                    metadata = {
                        "source": document["title"],
                        "source_name": document["title"],
                        "source_path": source_path,
                        "document_id": document["document_id"],
                        "source_title": document["title"],
                        "source_fingerprint": document["source_fingerprint"],
                        "chunk_index": int(c["chunk_index"]),
                        "chunk_type": c.get("chunk_type") or "body",
                        "vector_eligible": 1,
                        "embedding_model_id": c.get("embedding_model_id"),
                        "embedding_dimension": c.get("embedding_dimension"),
                        "embedding_normalized": c.get("embedding_normalized"),
                        "embedding_config_hash": c.get("embedding_config_hash"),
                        "parent_section_title": c.get("parent_section_title"),
                        "quality_flags": c.get("quality_flags"),
                        "ingested_at": document["last_ingested_at"],
                        "file_hash": current_hash,
                        "last_modified": document["last_modified"],
                        "file_size": int(document["file_size"] or 0),
                        "department": document["department"],
                        "application": document["application"],
                        "status": "active",
                        "extraction_method": c.get("extraction_method"),
                        "has_text_layer": int(c.get("has_text_layer", 1) or 0),
                        "ocr_confidence": c.get("ocr_confidence"),
                        "char_count": int(c.get("char_count") or 0) or None,
                        "word_count": int(c.get("word_count") or 0) or None,
                        "sheet_name": c.get("sheet_name"),
                        "range_ref": c.get("range_ref"),
                        "row_start": int(c["row_start"]) if c.get("row_start") is not None else None,
                        "row_end": int(c["row_end"]) if c.get("row_end") is not None else None,
                        "col_start": int(c["col_start"]) if c.get("col_start") is not None else None,
                        "col_end": int(c["col_end"]) if c.get("col_end") is not None else None,
                        "headers": json.dumps(c.get("headers"), ensure_ascii=False) if isinstance(c.get("headers"), list) else c.get("headers"),
                        "has_structured_preview": int(c.get("has_structured_preview", 0) or 0),
                        "file_type": document.get("file_type"),
                    }
                    if c.get("page_number") is not None:
                        metadata["page_number"] = int(c["page_number"])
                    if c.get("section_title"):
                        metadata["section_title"] = c["section_title"]
                    metadatas.append(_sanitize_chroma_metadata(metadata))

                collection.add(
                    ids=[c["chunk_id"] for c in vector_eligible_chunks],
                    embeddings=embeds,
                    documents=[c["text"] for c in vector_eligible_chunks],
                    metadatas=metadatas,
                )
                total_vectors += len(vector_eligible_chunks)

            insert_chunks(document["document_id"], chunks, document)
            ocr_chunk_count = sum(1 for chunk in chunks if chunk.get("extraction_method") == "ocr")
            if ocr_chunk_count:
                print(f"[Spark Ingest] Created {ocr_chunk_count} OCR-derived chunks from {filepath.name}")
            document["ingestion_status"] = "ingested"
            upsert_document(document)

            hashes[source_path] = current_hash
            ingested += 1
            total_chunks += len(chunks)
            print(f"[Spark Ingest]   {len(chunks)} chunks ingested ({len(vector_eligible_chunks)} vector-eligible).")

        if not target_source_path:
            removed_paths = purge_missing_documents(current_source_paths)
            for removed_path in removed_paths:
                _delete_chroma_source(collection, removed_path)
                hashes.pop(removed_path, None)
                print(f"[Spark Ingest]   Purged stale source from SQLite/Chroma/hash state: {removed_path}")

        save_hashes(hashes)
        finish_ingestion_run(
            run_id,
            status="completed",
            files_seen=processed,
            files_ingested=ingested,
            files_skipped=skipped,
            files_failed=failed,
            chunks_created=total_chunks,
            vector_chunks_created=total_vectors,
            vector_skipped_count=vector_skipped_run,
        )
        print(f"[Spark Ingest] Done. {total_chunks} chunks ({total_vectors} vectors) new. {skipped} file(s) unchanged.")
        return {
            "ok": True,
            "documents_processed": processed,
            "documents_indexed": ingested,
            "chunks_created": total_chunks,
            "sqlite_chunks_created": total_chunks,
            "vector_chunks_created": total_vectors,
            "vectors_created": total_vectors,
            "vector_skipped_count": vector_skipped_run,
            "failed_count": failed,
            "skipped_count": skipped,
            "failed": failed_items,
            "skipped": skipped_items,
            "unsupported": unsupported_items,
            "message": "Ingestion completed.",
        }
    except Exception as exc:
        finish_ingestion_run(
            run_id,
            status="failed",
            files_seen=len(files),
            files_ingested=ingested,
            files_skipped=skipped,
            files_failed=failed + 1,
            chunks_created=total_chunks,
            error=str(exc),
        )
        return {
            "ok": False,
            "documents_processed": processed,
            "documents_indexed": ingested,
            "chunks_created": total_chunks,
            "failed_count": failed + 1,
            "skipped_count": skipped,
            "failed": failed_items + [{"source_path": target_source_path or "*", "reason": str(exc)}],
            "skipped": skipped_items,
            "unsupported": unsupported_items,
            "message": f"Ingestion failed: {exc}",
        }

if __name__ == "__main__":
    import argparse

    def _purge_collection_vectors() -> int:
        collection = _get_chroma_collection()
        before = int(collection.count() or 0)
        if before <= 0:
            return 0
        try:
            collection.delete(where={})
            return before
        except Exception:
            # Fallback for Chroma variants that do not accept empty where.
            deleted = 0
            batch_size = 500
            while True:
                rows = collection.get(limit=batch_size, offset=0, include=[])
                ids = rows.get("ids") or []
                if not ids:
                    break
                collection.delete(ids=ids)
                deleted += len(ids)
                if len(ids) < batch_size:
                    break
            return deleted

    parser = argparse.ArgumentParser(description="Spark ingestion (SQLite + Chroma vectors)")
    parser.add_argument("--source-path", default=None, help="Relative intake source path to ingest (optional)")
    parser.add_argument("--force-reingest", action="store_true", help="Re-ingest even if unchanged")
    parser.add_argument("--purge-vectors", action="store_true", help="Delete all vectors from the active Chroma collection")
    parser.add_argument("--purge-hashes", action="store_true", help="Delete the per-collection hash store file")
    parser.add_argument(
        "--rebuild",
        action="store_true",
        help="Shortcut: purge vectors + purge hashes + force re-ingest (active collection only)",
    )
    args = parser.parse_args()

    if args.rebuild:
        args.purge_vectors = True
        args.purge_hashes = True
        args.force_reingest = True

    if args.purge_vectors:
        deleted = _purge_collection_vectors()
        print(f"[Spark Ingest] Purged vectors: {deleted} collection={CHROMA_COLLECTION}")

    if args.purge_hashes:
        try:
            if HASH_STORE.exists():
                HASH_STORE.unlink()
                print(f"[Spark Ingest] Deleted hash store: {HASH_STORE}")
        except Exception as exc:
            raise RuntimeError(f"Failed to delete hash store {HASH_STORE}: {exc}") from exc

    run_ingestion(target_source_path=args.source_path, force_reingest=args.force_reingest)
